from pathlib import Path
from fastapi.responses import FileResponse
import os
import asyncio
from docx.shared import Inches
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Request, Depends, HTTPException
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import selectinload

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.future import select

from database import get_db
from services.gerar_pdf_service import gerar_pdf


# MODELOS
from models.turma import Turma
from models.aluno import Aluno
from models.escola import Escola
from models.professor import Professor
from models.matricula import Matricula
from models.distribuicao import Distribuicao


router = APIRouter(
    prefix="/pdf",
    tags=["PDF"]
)


templates = Jinja2Templates(
    directory="templates"
)

DOWNLOAD_DIR = Path("static/downloads")
DOWNLOAD_DIR.mkdir(parents=True, exist_ok=True)

def proximo_nome(prefixo: str, extensao: str):

    numero = 1

    while (DOWNLOAD_DIR / f"{prefixo}_{numero}.{extensao}").exists():
        numero += 1

    return f"{prefixo}_{numero}.{extensao}"

@router.get("/turma/{turma_id}")
async def gerar_pdf_turma(

    request: Request,

    turma_id:int,

    db: AsyncSession = Depends(get_db)

):


    # ===============================
    # BUSCAR TURMA
    # ===============================

    result = await db.execute(

        select(Turma)
        .options(
            selectinload(Turma.classe)
        )
        .where(
            Turma.id == turma_id
        )

    )

    turma = result.scalar_one_or_none()



    if not turma:

        raise HTTPException(
            status_code=404,
            detail="Turma não encontrada"
        )



    # ===============================
    # BUSCAR ALUNOS PELA MATRICULA
    # ===============================


    result = await db.execute(

        select(Aluno)
        .join(
            Matricula,
            Matricula.aluno_id == Aluno.id
        )
        .where(
            Matricula.turma_id == turma_id
        )
        .order_by(
            Aluno.nome
        )

    )


    alunos = result.scalars().all()

    # ===============================
    # BUSCAR ESCOLA
    # ===============================

    result_escola = await db.execute(
        select(Escola)
    )

    escola = result_escola.scalar_one_or_none()

    # ===============================
    # BUSCAR CLASSE
    # ===============================

    classe = turma.classe

    # ===============================
    # BUSCAR PROFESSOR
    # ===============================

    result_prof = await db.execute(

        select(Professor)
        .join(
            Distribuicao,
            Distribuicao.professor_id == Professor.id
        )
        .where(
            Distribuicao.turma_id == turma_id
        )

    )

    professor = result_prof.scalar_one_or_none()



    homens = 0
    mulheres = 0



    for aluno in alunos:


        if aluno.sexo:

            sexo = aluno.sexo.upper()


            if sexo in ["M","H"]:

                homens += 1


            elif sexo == "F":

                mulheres += 1



    total = homens + mulheres


    ano = 2026




    # ===============================
    # CONTEXTO
    # ===============================

    contexto = {

        "request": request,

        "escola": escola,

        "classe": classe,

        "turma": turma,

        "alunos": alunos,

        "professor": professor,

        "homens": homens,

        "mulheres": mulheres,

        "total": total,

        "ano": ano
    }

    # ===============================
    # GERAR HTML
    # ===============================

    html = templates.get_template(
        "professor_turma.html"
    ).render(contexto)

    pdf = await asyncio.to_thread(
        gerar_pdf,
        html
    )

    nome_pdf = proximo_nome(
        "lista_turma",
        "pdf"
    )

    caminho = DOWNLOAD_DIR / nome_pdf

    with open(caminho, "wb") as f:
        f.write(pdf)

    return FileResponse(
        path=caminho,
        filename=nome_pdf,
        media_type="application/pdf"
    )

@router.get("/word/turma/{turma_id}")
async def gerar_word_turma(

    request: Request,
    turma_id: int,
    db: AsyncSession = Depends(get_db)

):

    # ===============================
    # BUSCAR TURMA
    # ===============================

    result = await db.execute(

        select(Turma)
        .options(selectinload(Turma.classe))
        .where(Turma.id == turma_id)

    )

    turma = result.scalar_one_or_none()

    if not turma:
        raise HTTPException(
            status_code=404,
            detail="Turma não encontrada"
        )

    # ===============================
    # BUSCAR ALUNOS
    # ===============================

    result = await db.execute(

        select(Aluno)
        .join(
            Matricula,
            Matricula.aluno_id == Aluno.id
        )
        .where(
            Matricula.turma_id == turma_id
        )
        .order_by(Aluno.nome)

    )

    alunos = result.scalars().all()

    # ===============================
    # ESCOLA
    # ===============================

    result = await db.execute(
        select(Escola)
    )

    escola = result.scalar_one_or_none()

    # ===============================
    # PROFESSOR
    # ===============================

    result = await db.execute(

        select(Professor)
        .join(
            Distribuicao,
            Distribuicao.professor_id == Professor.id
        )
        .where(
            Distribuicao.turma_id == turma_id
        )

    )

    professor = result.scalar_one_or_none()

    classe = turma.classe

    homens = sum(1 for a in alunos if a.sexo and a.sexo.upper() in ["M", "H"])
    mulheres = sum(1 for a in alunos if a.sexo and a.sexo.upper() == "F")

    # ===============================
    # DOCUMENTO WORD
    # ===============================

    doc = Document()

    secao = doc.sections[0]
    secao.top_margin = Pt(28)
    secao.bottom_margin = Pt(28)
    secao.left_margin = Pt(28)
    secao.right_margin = Pt(28)

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ===============================
    # INSERIR LOGOTIPO
    # ===============================

    imagem = os.path.join(
        os.getcwd(),
        "static",
        "img",
        "img.png"
    )

    try:

        run = titulo.add_run()

        run.add_picture(
            imagem,
            width=Inches(1)
        )

    except Exception as e:

        print("Erro ao carregar imagem:", e)

    titulo.add_run("\nREPÚBLICA DE MOÇAMBIQUE\n").bold = True
    titulo.add_run("Ministério da Educação e Cultura\n")

    if escola:
        titulo.add_run(f"{escola.provincia}\n")
        titulo.add_run(f"{escola.distrito}\n")
        titulo.add_run(f"{escola.nome}\n")

    p = doc.add_paragraph()

    p.add_run(f"Classe: {classe.classe}      ")
    p.add_run(f"Turma: {turma.turma}      ")
    p.add_run("Ano: 2026")

    t = doc.add_heading(
        "LISTA NOMINAL DOS ALUNOS",
        level=2
    )

    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tabela = doc.add_table(
        rows=1,
        cols=5
    )

    tabela.style = "Table Grid"

    cab = tabela.rows[0].cells

    cab[0].text = "Nº"
    cab[1].text = "Nome Completo"
    cab[2].text = "Sexo"
    cab[3].text = "Data Nascimento"
    cab[4].text = "Estado"

    for i, aluno in enumerate(alunos, start=1):

        linha = tabela.add_row().cells

        linha[0].text = str(i)
        linha[1].text = aluno.nome
        linha[2].text = aluno.sexo or ""

        linha[3].text = (
            aluno.data_nascimento.strftime("%d/%m/%Y")
            if aluno.data_nascimento else ""
        )

        linha[4].text = "Ativo"

    doc.add_paragraph()

    doc.add_heading("ESTATÍSTICA", level=3)

    doc.add_paragraph(f"Homens (H): {homens}")
    doc.add_paragraph(f"Mulheres (M): {mulheres}")
    doc.add_paragraph(f"HM: {homens + mulheres}")

    doc.add_paragraph("\n\n")

    assinatura = doc.add_paragraph()
    assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER

    assinatura.add_run("O Director da Turma\n\n\n")
    assinatura.add_run("_________________________\n")

    if professor:
        assinatura.add_run(f"({professor.nome})")

    arquivo = BytesIO()

    doc.save(arquivo)

    arquivo.seek(0)

    nome_word = proximo_nome(
        "lista_turma",
        "docx"
    )

    caminho = DOWNLOAD_DIR / nome_word

    with open(caminho, "wb") as f:
        f.write(arquivo.read())

    return FileResponse(
        path=caminho,
        filename=nome_word,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
